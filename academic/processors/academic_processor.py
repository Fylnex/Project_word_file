from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


class AcademicProcessor:
    def __init__(self):
        self.document = Document()
        self.image_counter = 1
        self.code_counter = 1

        # Настройка стилей документа
        self._set_styles()

    def _set_styles(self):
        # Установка шрифта по умолчанию
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Создание стиля для заголовка статьи
        title_style = self.document.styles.add_style('ArticleTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(16)
        title_font.bold = True

        # Создание стиля для авторов
        author_style = self.document.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
        author_font = author_style.font
        author_font.name = 'Times New Roman'
        author_font.size = Pt(14)

        # Создание стиля для аннотации
        abstract_style = self.document.styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH)
        abstract_font = abstract_style.font
        abstract_font.name = 'Times New Roman'
        abstract_font.size = Pt(12)
        abstract_font.italic = True

        # Создание стиля для заголовков первого уровня
        heading1_style = self.document.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Times New Roman'
        heading1_font.size = Pt(14)
        heading1_font.bold = True

        # Создание стиля для заголовков второго уровня
        heading2_style = self.document.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'Times New Roman'
        heading2_font.size = Pt(12)
        heading2_font.bold = True

    def add_title(self, title, authors, abstract, keywords):
        # Проверка наличия необходимых данных
        if not title or not authors or not abstract or not keywords:
            print("Ошибка: Недостаточно данных для добавления заголовка статьи.")
            return

        # Добавление названия статьи
        p = self.document.add_paragraph(title, style='ArticleTitle')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавление авторов
        p = self.document.add_paragraph(authors, style='Author')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавление аннотации
        self.document.add_heading('Аннотация', level=1)
        self.document.add_paragraph(abstract, style='Abstract')

        # Добавление ключевых слов
        self.document.add_heading('Ключевые слова', level=1)
        self.document.add_paragraph(', '.join(keywords))

    def add_paragraph(self, text):
        if not text:
            print("Предупреждение: Пустой текстовый параграф не был добавлен.")
            return
        self.document.add_paragraph(text)

    def add_heading(self, heading):
        if not heading:
            print("Предупреждение: Пустой заголовок не был добавлен.")
            return
        self.document.add_heading(heading, level=1)

    def add_subheading(self, subheading):
        if not subheading:
            print("Предупреждение: Пустой подзаголовок не был добавлен.")
            return
        self.document.add_heading(subheading, level=2)

    def add_image(self, image_path, caption):
        if not os.path.exists(image_path):
            print(f"Ошибка: Файл изображения '{image_path}' не найден. Блок изображения не был добавлен.")
            return
        try:
            self.document.add_picture(image_path, width=Inches(5))
            last_paragraph = self.document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавление подписи к изображению
            caption_text = f'Рисунок №{self.image_counter} – {caption}'
            p = self.document.add_paragraph(caption_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.image_counter += 1
        except Exception as e:
            print(f"Ошибка при добавлении изображения: {e}. Блок изображения не был добавлен.")

    def add_numbered_list(self, items):
        if not items:
            print("Предупреждение: Пустой нумерованный список не был добавлен.")
            return
        for item in items:
            self.document.add_paragraph(item, style='List Number')

    def add_unordered_list(self, items):
        if not items:
            print("Предупреждение: Пустой ненумерованный список не был добавлен.")
            return
        for item in items:
            self.document.add_paragraph(item, style='List Bullet')

    def add_code_block(self, code_title, code_content):
        if not code_content:
            print("Предупреждение: Пустой блок кода не был добавлен.")
            return
        try:
            # Добавление названия блока кода
            caption_text = f'Листинг №{self.code_counter} – {code_title}'
            self.document.add_paragraph(caption_text, style='Normal')

            # Добавление самого кода
            code_paragraph = self.document.add_paragraph(style='Normal')
            code_paragraph.paragraph_format.left_indent = Inches(0.5)
            code_run = code_paragraph.add_run(code_content)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)

            self.code_counter += 1
        except Exception as e:
            print(f"Ошибка при добавлении блока кода: {e}. Блок кода не был добавлен.")

    def add_table(self, table_data):
        if not table_data:
            print("Предупреждение: Пустая таблица не была добавлена.")
            return
        # Пока заготовка для таблицы
        pass

    def add_bibliography(self, references):
        if not references:
            print("Предупреждение: Пустой список литературы не был добавлен.")
            return
        self.document.add_heading('Список литературы', level=1)
        for ref in references:
            self.document.add_paragraph(ref, style='List Number')

    def save_to_word(self, file_path):
        try:
            name_file_word = os.path.splitext(os.path.basename(file_path))[0] + '.docx'
            download_path = os.path.join(os.path.expanduser('~'), 'Downloads')
            output_path = os.path.join(download_path, f"{name_file_word}")
            self.document.save(output_path)
            print(f"Документ успешно сохранен по пути: {file_path}")
        except Exception as e:
            print(f"Ошибка при сохранении документа: {e}")
