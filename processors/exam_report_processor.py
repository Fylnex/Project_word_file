from .document_processor import DocumentProcessor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx import Document

class ExamReportProcessor(DocumentProcessor):
    """
    Класс для обработки отчета по экзамену.
    """
    def __init__(self, template_path, output_path, replacements, students_data):
        super().__init__(template_path, output_path, replacements)
        self.students_data = students_data

    def process(self):
        """Процесс обработки отчета по экзамену."""
        try:
            self.load_template()
            self.replace_flags()
            self.insert_table()
            self.save_document()
            return True
        except Exception as e:
            print(f"Ошибка при обработке документа: {e}")
            return False

    def replace_flags(self):
        """Замена флагов в документе."""
        for paragraph in self.doc.paragraphs:
            for flag, value in self.replacements.items():
                if flag in paragraph.text:
                    paragraph.text = paragraph.text.replace(flag, str(value))

        # Поиск таблицы
        # if len(self.doc.tables):
        #     print("есть таблица")
        # else:print("нет таблицы")




    def insert_table(self):
        """Вставка таблицы с данными студентов на место флага ${table}."""
        for i, paragraph in enumerate(self.doc.paragraphs):
            if "${table}" in paragraph.text:
                p = paragraph._element
                # Создаем таблицу и вставляем ее после удаления параграфа
                tbl = self.create_table_with_borders()
                tbl_element = tbl._element
                p.addnext(tbl_element)
                p.getparent().remove(p)
                break

    def create_table_with_borders(self):
        """Создание таблицы с границами и правильным форматированием столбцов."""
        # Создаем таблицу с двумя строками для заголовков и количеством столбцов 7
        table = self.doc.add_table(rows=2, cols=7)

        # Устанавливаем стиль таблицы для добавления границ ко всем ячейкам
        # table.style = 'Table Grid'

        # Объединяем ячейки для заголовка "Экзаменационная оценка"
        table.cell(0, 4).merge(table.cell(0,5))

        # Объединяем заголовки с ячейками во втором ряду для других столбцов
        table.cell(0, 0).merge(table.cell(1, 0))  # №
        table.cell(0, 1).merge(table.cell(1, 1))  # Фамилия и инициалы
        table.cell(0, 2).merge(table.cell(1, 2))  # Группа
        table.cell(0, 3).merge(table.cell(1, 3))  # № зач. книжки
        table.cell(0, 6).merge(table.cell(1, 6))  # Подпись экзаменатора



        # заголовки для второй строки
        subheaders = ['№', 'Фамилия и инициалы', 'Группа', '№ зач. книжки', 'цифрой', 'прописью', 'Подпись экзаменатора']
        for i, subheader in enumerate(subheaders):
            table.cell(1, i).text = subheader

        # заголовки для первой строки после объединения
        headers = ['Экзаменационная оценка',]
        table.cell(0,4).text = headers


        # Добавляем строки для каждого студента
        for student in self.students_data:
            row = table.add_row().cells
            row[0].text = student.get('№', '')
            row[1].text = student.get('Фамилия и инициалы', '')
            row[2].text = student.get('Группа', '')
            row[3].text = student.get('№ зач. книжки', '')
            row[4].text = student.get('Экзаменационная оценка цифрой', '')
            row[5].text = student.get('Экзаменационная оценка прописью', '')
            row[6].text = student.get('Подпись экзаменатора', '')




        # Устанавливаем ширину столбцов
        widths = [0.1, 2, 0.4, 0.4, 1, 1, 1]  # Пример ширины для каждого столбца в дюймах
        for i, col in enumerate(table.columns):
            for cell in col.cells:
                cell.width = Inches(widths[i])

        # Добавляем границы к таблице
        # table.style = 'Table Grid'
        self.add_table_borders(table)

        return table

    def add_table_borders(self, table):
        """Добавление границ к таблице."""
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')

                # Создаем границы для каждой стороны ячейки
                borders = {
                    'top': 'single',
                    'left': 'single',
                    'bottom': 'single',
                    'right': 'single'
                }

                for border_name, border_val in borders.items():
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), border_val)
                    border.set(qn('w:sz'), '4')  # Толщина границы
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)

                tcPr.append(tcBorders)