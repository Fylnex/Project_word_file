from .document_processor import DocumentProcessor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches



class CouncilReportProcessor(DocumentProcessor):
    """
    Класс для обработки отчета по научному совету.
    """
    def __init__(self, template_path, output_path, replacements, employees_data):
        super().__init__(template_path, output_path, replacements)
        self.employees_data = employees_data

    def process(self):
        """Процесс обработки отчета по научному совету."""
        try:
            self.load_template()
            self.replace_flags()
            self.insert_table()
            self.save_document()
            return True
        except Exception as e:
            print(f"Ошибка при обработке документа: {e}")
            return False

    def replace_flags(self):
        """Замена флагов в документе, включая таблицы."""
        # Проход по всем параграфам документа
        for paragraph in self.doc.paragraphs:
            for flag, value in self.replacements.items():
                if flag in paragraph.text:
                    paragraph.text = paragraph.text.replace(flag, str(value))

        # Проход по всем таблицам в документе
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Проход по всем параграфам в каждой ячейке
                    for paragraph in cell.paragraphs:
                        for flag, value in self.replacements.items():
                            if flag in paragraph.text:
                                paragraph.text = paragraph.text.replace(flag, str(value))


    def insert_table(self):
        """Вставка таблицы с данными сотрудников на место флага ${table}."""
        for i, paragraph in enumerate(self.doc.paragraphs):
            if "${table}" in paragraph.text:
                p = paragraph._element
                # Создаем таблицу и вставляем ее после удаления параграфа
                tbl = self.create_table_with_borders()
                tbl_element = tbl._element
                p.addnext(tbl_element)
                p.getparent().remove(p)
                break

    def create_table_with_borders(self):
        """Создание таблицы"""
        # Создаем таблицу с 1 строкой для заголовков и количеством столбцов 6
        table = self.doc.add_table(rows=1, cols=6)

        # Устанавливаем границы таблицы
        self.add_table_borders(table)

        # Заголовки для второй строки
        subheaders = ['№№ ПП', 'ФАМИЛИЯ, И., О.', 'УЧЕНАЯ СТЕПЕНЬ, УЧЕНОЕ ЗВАНИЕ', 'ДОЛЖНОСТЬ', 'РАСПИСКА В ЯВКЕ НА ЗАСЕДАНИЕ СОВЕТА', 'РАСПИСКА В ПОЛУЧЕНИИ БЮЛЛЕТЕНЕЙ']
        for i, subheader in enumerate(subheaders):
            table.cell(0, i).text = subheader



        # Добавляем строки для каждого студента
        n=0
        for emloyee in self.employees_data:
            row = table.add_row().cells
            n+=1
            row[0].text = str(n)
            row[1].text = emloyee.get('ФАМИЛИЯ, И., О.', '')
            row[2].text = emloyee.get('УЧЕНАЯ СТЕПЕНЬ, УЧЕНОЕ ЗВАНИЕ', '')
            row[3].text = emloyee.get('ДОЛЖНОСТЬ', '')
            row[4].text = ''
            row[5].text = ''

        # Устанавливаем ширину столбцов
        widths = [0.1, 4, 0.5, 0.5, 1, 1]  # Устанавливаем ширину для каждого столбца в дюймах
        for i, col in enumerate(table.columns):
            for cell in col.cells:
                cell.width = Inches(widths[i])

        # Добавляем границы к таблице
        self.add_table_borders(table)

        return table

    def add_table_borders(self, table):
        # Добавление границ ко всем ячейкам таблицы
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')

                # Установка границ для каждой стороны ячейки
                borders = ['top', 'left', 'bottom', 'right']
                for border_name in borders:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')  # Тип линии
                    border.set(qn('w:sz'), '4')  # Толщина границы
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)

                tcPr.append(tcBorders)
